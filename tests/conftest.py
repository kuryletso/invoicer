from typing import Any, Callable
from decimal import Decimal
from pathlib import Path

import pytest
from docx import Document
from sqlalchemy import create_engine
from sqlalchemy.orm import Session

from app.core.diagnostics import DiagnosticCollector
from app.document_engine.blueprint.models.template import TemplateConfig
from app.document_engine.blueprint.template_builder import TemplateBuilderContext
from app.document_engine.normalization.models.inlines import (
    NormalizedTextNode,
    NormalizedTextStyle,
)
from app.document_engine.enums.enums import PlaceholderType, MoneySymbolPosition

from app.db.base import Base
from app.db.models import *  # noqa: F401,F403  — registers every mapper
from app.db.models.configs.default_template_config import DefaultTemplateConfig
from app.db.models.core.bank_account import BankAccount
from app.db.models.core.bank_account_localization import BankAccountLocalization
from app.db.models.core.document_sequence import DocumentSequence
from app.db.models.core.invoice_line import InvoiceLine
from app.db.models.core.invoice_line_localization import InvoiceLineLocalization
from app.db.models.core.organization import Organization
from app.db.models.core.organization_localization import OrganizationLocalization
from app.db.models.core.representative import Representative
from app.db.models.core.representative_localization import RepresentativeLocalization
from app.db.models.core.tax_id import TaxId
from app.db.models.references.country import Country
from app.db.models.references.country_localization import CountryLocalization
from app.db.models.references.currency import Currency
from app.db.models.references.currency_localization import CurrencyLocalization
from app.db.models.references.language import Language
from app.db.models.registries.document_type import DocumentTypeRegistry
from app.db.models.registries.measurement_unit import MeasurementUnitRegistry
from app.db.models.registries.measurement_unit_localization import (
    MeasurementUnitRegistryLocalization,
)
from app.db.models.registries.placeholder import PlaceholderRegistry
from app.db.models.registries.placeholder_localization import (
    PlaceholderRegistryLocalization,
)
from app.db.models.registries.tax_id_system import TaxIdSystemRegistry
from app.db.models.registries.tax_id_system_localization import (
    TaxIdSystemRegistryLocalization,
)

from app.services.language import LanguageSpec


LANGS = (LanguageSpec(code="UKR", alpha_2="uk"), LanguageSpec(code="ENG", alpha_2="en"))
CODES = tuple(lang.code for lang in LANGS)


class FixtureInputProvider:
    """In-memory TemplateInputProvider for tests/dev — no database required.

    Structurally satisfies orchestration.ports.TemplateInputProvider.
    """

    def __init__(
        self,
        *,
        languages: set[str] | None = None,
        placeholders: dict[str, dict[str, Any]] | None = None,
        config: TemplateConfig | None = None,
    ) -> None:
        self._languages = languages
        self._placeholders = placeholders or {
            "org_name": {"active": True, "required": True, "type": PlaceholderType.SCALAR},
            "invoice_no": {"active": True, "required": True, "type": PlaceholderType.SCALAR},
        }
        self._config = config or TemplateConfig(
            primary_language="ENG",
            secondary_language="UKR",
            type="invoice",
            name="seed",
            description="",
            append_currency=False,
        )

    def default_template_config(self) -> TemplateConfig:
        return self._config

    def placeholder_defaults(self) -> dict[str, dict[str, Any]]:
        return self._placeholders

    def languages(self) -> set[str]:
        """Mirrors DbTemplateInputProvider: the template's renderable languages."""
        if self._languages is not None:
            return self._languages
        return {
            code
            for code in (self._config.primary_language, self._config.secondary_language)
            if code
        }


@pytest.fixture
def fixture_provider() -> FixtureInputProvider:
    return FixtureInputProvider()


@pytest.fixture
def diagnostics() -> DiagnosticCollector:
    return DiagnosticCollector()


@pytest.fixture
def make_docx(tmp_path: Path) -> Callable[..., Path]:
    """Factory building a real .docx via python-docx into the test's tmp dir."""

    def _make(
        paragraphs: list[str] | tuple[str, ...] = (),
        table: list[list[str]] | None = None,
        name: str = "test.docx",
    ) -> Path:
        document = Document()
        for text in paragraphs:
            document.add_paragraph(text)

        if table:
            rows, cols = len(table), len(table[0])
            built = document.add_table(rows=rows, cols=cols)
            for r, row in enumerate(table):
                for c, value in enumerate(row):
                    built.cell(r, c).text = value

        path = tmp_path / name
        document.save(path)
        return path

    return _make


@pytest.fixture
def make_text_node() -> Callable[[str], NormalizedTextNode]:
    def _make(text: str) -> NormalizedTextNode:
        return NormalizedTextNode(
            text=text,
            style=NormalizedTextStyle(
                bold=False,
                italic=False,
                underline=False,
                font_name="Calibri",
                font_size=22,
                color="000000",
            ),
        )

    return _make


@pytest.fixture
def session() -> Session:
    """In-memory SQLite with the full schema and seeded reference data."""

    engine = create_engine("sqlite://")
    Base.metadata.create_all(engine)

    with Session(engine) as s:
        s.add_all([
            Language(code="ENG", code_alpha_2="en", label_en="English", label_uk="Англійська"),
            Language(code="UKR", code_alpha_2="uk", label_en="Ukrainian", label_uk="Українська"),
            DocumentTypeRegistry(code="invoice"),
        ])

        s.add(Currency(
            code="UAH",
            decimal_places=2,
            decimal_separator=",",
            grouping_separator=" ",
            symbol_position=MoneySymbolPosition.SUFFIX,
            symbol_spacing=True,
            localizations={
                "UKR": CurrencyLocalization(language_code="UKR", name="Гривня", symbol="₴"),
                "ENG": CurrencyLocalization(language_code="ENG", name="Hryvnia", symbol="UAH"),
            },
        ))

        s.add(Country(
            code="UKR",
            localizations={
                "UKR": CountryLocalization(language_code="UKR", name="Україна"),
                "ENG": CountryLocalization(language_code="ENG", name="Ukraine"),
            },
        ))

        s.add(TaxIdSystemRegistry(
            code="edrpou",
            localizations={
                "UKR": TaxIdSystemRegistryLocalization(language_code="UKR", name="ЄДРПОУ"),
                "ENG": TaxIdSystemRegistryLocalization(language_code="ENG", name="EDRPOU"),
            },
        ))

        s.add(MeasurementUnitRegistry(
            code="hour",
            localizations={
                "UKR": MeasurementUnitRegistryLocalization(language_code="UKR", name="год"),
                "ENG": MeasurementUnitRegistryLocalization(language_code="ENG", name="hr"),
            },
        ))

        s.commit()
        yield s


@pytest.fixture
def make_org(session: Session) -> Callable[..., Organization]:
    """Factory for an Organization with localizations, tax id, rep and bank account."""

    def _make(
        name: str = "Acme",
        *,
        email: str | None = "acme@example.com",
        phone: str | None = None,
        with_representative: bool = True,
        with_bank: bool = True,
        tax_value: str = "12345678",
    ) -> Organization:
        org = Organization(
            email=email,
            phone=phone,
            localizations={
                "UKR": OrganizationLocalization(
                    language_code="UKR", org_type="ТОВ",
                    legal_name=f"{name} UA", address="вул. Головна, 1",
                ),
                "ENG": OrganizationLocalization(
                    language_code="ENG", org_type="LLC",
                    legal_name=name, address="1 Main St",
                ),
            },
        )
        org.tax_ids.append(
            TaxId(tax_id_system_code="edrpou", country_code="UKR", value=tax_value)
        )

        if with_representative:
            org.representatives.append(Representative(
                localizations={
                    "UKR": RepresentativeLocalization(
                        language_code="UKR", name="Іван Петренко", title="Директор",
                    ),
                    "ENG": RepresentativeLocalization(
                        language_code="ENG", name="Ivan Petrenko", title="Director",
                    ),
                },
            ))

        if with_bank:
            org.bank_accounts.append(BankAccount(
                country_code="UKR",
                currency_code="UAH",
                iban=f"UA00{tax_value}0000000000000",
                swift="TESTUA00",
                localizations={
                    "UKR": BankAccountLocalization(
                        language_code="UKR", bank_name="ПриватБанк", bank_info="МФО 305299",
                    ),
                    "ENG": BankAccountLocalization(
                        language_code="ENG", bank_name="PrivatBank", bank_info="MFO 305299",
                    ),
                },
            ))

        session.add(org)
        session.commit()
        return org

    return _make


@pytest.fixture
def make_line(session: Session) -> Callable[..., InvoiceLine]:
    """Factory for a stored InvoiceLine."""

    def _make(
        description: str = "Design work",
        *,
        quantity: str = "10.000",
        unit_price: str = "125.50",
        tax_rate: str = "0.20000",
    ) -> InvoiceLine:
        line = InvoiceLine(
            quantity=Decimal(quantity),
            unit_price=Decimal(unit_price),
            tax_rate=Decimal(tax_rate),
            measurement_unit_code="hour",
            localizations={
                "UKR": InvoiceLineLocalization(language_code="UKR", description="Дизайн"),
                "ENG": InvoiceLineLocalization(language_code="ENG", description=description),
            },
        )
        session.add(line)
        session.commit()
        return line

    return _make


@pytest.fixture
def make_sequence(session: Session) -> Callable[..., DocumentSequence]:
    """Factory for a DocumentSequence bound to an organization."""

    def _make(
        organization: Organization,
        *,
        prefix: str = "INV-",
        counter: int = 6,
        padding: int = 4,
    ) -> DocumentSequence:
        sequence = DocumentSequence(
            document_type_code="invoice",
            organization_id=organization.id,
            prefix=prefix,
            counter=counter,
            padding=padding,
        )
        session.add(sequence)
        session.commit()
        return sequence

    return _make


@pytest.fixture
def make_context() -> Callable[..., TemplateBuilderContext]:
    """Factory for a TemplateBuilderContext with its own DiagnosticCollector.

    Inspect emitted warnings afterwards via ``context.diagnostics.warnings``.
    """

    def _make(
        *,
        default_language: str = "ENG",
        languages: set[str] | None = None,
        placeholders: dict[str, dict[str, Any]] | None = None,
        diagnostics: DiagnosticCollector | None = None,
    ) -> TemplateBuilderContext:
        return TemplateBuilderContext(
            default_language=default_language,
            placeholder_defaults=placeholders
            or {
                "org_name": {"active": True, "required": True, "type": PlaceholderType.SCALAR},
                "invoice_no": {"active": True, "required": True, "type": PlaceholderType.SCALAR},
            },
            languages=languages or {"ENG", "UKR"},
            placeholders={},
            diagnostics=diagnostics or DiagnosticCollector(),
        )

    return _make
