"""Tests for the rendering stage: resolver (blueprint + context -> ResolvedDocument)
and the docx emitter (ResolvedDocument -> .docx bytes, re-opened with python-docx)."""
import base64
import zipfile
from io import BytesIO

import pytest
from docx import Document

from app.core.diagnostics import DiagnosticCollector

from app.document_engine.enums.enums import (
    ParagraphAlignment, SectionType, PageOrientation, VerticalAlignment,
    TableCellShading, TableBorderStyleEnum, TableWidthType, PlaceholderType,
    HeaderFooterType,
)
from app.document_engine.blueprint.models.segment import (
    TextStyleBlueprint, TextSegment, PlaceholderSegment,
    JoinedPlaceholderSegment, ImageSegment,
)
from app.document_engine.blueprint.models.paragraph import (
    ParagraphStyleBlueprint, ParagraphBlueprint,
)
from app.document_engine.blueprint.models.margins import MarginsBlueprint
from app.document_engine.blueprint.models.table import (
    TableStyleBlueprint, TableWidthBlueprint, TableBorderBlueprint,
    RowStyleBlueprint, CellStyleBlueprint,
    TableBlueprint, RowBlueprint, CellBlueprint,
    RowPlaceholder, CellPlaceholder, TablePlaceholder,
)
from app.document_engine.blueprint.models.header_footer import (
    HeaderFooterBlueprint, HeaderFooterGroupBlueprint,
)
from app.document_engine.blueprint.models.section import (
    SectionStyleBlueprint, SectionBlueprint,
)
from app.document_engine.blueprint.models.template import (
    TemplateBlueprint, TemplateConfig, PlaceholderDefinition,
)

from app.document_engine.rendering.context import (
    RenderContext, InvoiceTableData, InvoiceLineRow,
)
from app.document_engine.rendering.ports import Asset
from app.document_engine.rendering.errors import PlaceholderError
from app.document_engine.rendering.resolve.resolver import DocumentResolver
from app.document_engine.rendering.resolve.models import (
    ResolvedTextRun, ResolvedImageRun,
)
from app.document_engine.rendering.docx.emitter import DocxEmitter


# --------------------------------------------------------------------------
# blueprint factories
# --------------------------------------------------------------------------

def text_style(**over) -> TextStyleBlueprint:
    return TextStyleBlueprint(**{
        "bold": False, "italic": False, "underline": False,
        "font_name": "Calibri", "font_size": 22, "color": "000000", **over,
    })


def para_style(**over) -> ParagraphStyleBlueprint:
    return ParagraphStyleBlueprint(**{
        "alignment": ParagraphAlignment.LEFT, "spacing_before": 0, "spacing_after": 0,
        "indent_left": 0, "indent_right": 0, "keep_next": False, **over,
    })


def text_seg(text: str) -> TextSegment:
    return TextSegment(type="text", text=text, style=text_style())


def ph_seg(key: str, language: str = "ENG", ph_type: PlaceholderType = PlaceholderType.SCALAR) -> PlaceholderSegment:
    return PlaceholderSegment(type="placeholder", key=key, language=language,
                              ph_type=ph_type, style=text_style())


def paragraph(*segments) -> ParagraphBlueprint:
    return ParagraphBlueprint(type="paragraph", segments=tuple(segments), style=para_style())


def _margins() -> MarginsBlueprint:
    return MarginsBlueprint(top=0, bottom=0, left=108, right=108)


def cell_style(grid_span: int = 1) -> CellStyleBlueprint:
    return CellStyleBlueprint(shading=TableCellShading.CLEAR, shading_fill="auto",
                              margins=_margins(), grid_span=grid_span,
                              v_alignment=VerticalAlignment.CENTER)


def table_style() -> TableStyleBlueprint:
    b = TableBorderBlueprint(style=TableBorderStyleEnum.SINGLE, size=4, color="000000")
    return TableStyleBlueprint(
        width=TableWidthBlueprint(value=9000, type=TableWidthType.DXA), autofit=False,
        border_top=b, border_bottom=b, border_left=b, border_right=b,
        border_inside_v=b, border_inside_h=b, margins=_margins(),
    )


def static_cell(text: str) -> CellBlueprint:
    return CellBlueprint(type="cell", blocks=(paragraph(text_seg(text)),), style=cell_style())


def column_cell(key: str, language: str = "ENG") -> CellPlaceholder:
    return CellPlaceholder(type="placeholder_cell", key=key, language=language,
                           cell_style=cell_style(), text_style=text_style(), para_style=para_style())


def section_style() -> SectionStyleBlueprint:
    return SectionStyleBlueprint(
        section_type=SectionType.NEXTPAGE, page_width=11906, page_height=16838,
        orientation=PageOrientation.PORTRAIT, margin_header=708, margin_footer=708,
        margins=MarginsBlueprint(top=1440, bottom=1440, left=1440, right=1440),
    )


def empty_group() -> HeaderFooterGroupBlueprint:
    return HeaderFooterGroupBlueprint(default=None, first=None, even=None)


def section(blocks, headers=None, footers=None) -> SectionBlueprint:
    return SectionBlueprint(blocks=tuple(blocks), headers=headers or empty_group(),
                            footers=footers or empty_group(), style=section_style())


def config() -> TemplateConfig:
    return TemplateConfig(primary_language="ENG", secondary_language=None,
                          type="invoice", name="t", description="", append_currency=False)


def blueprint(sections, placeholders=None) -> TemplateBlueprint:
    return TemplateBlueprint(sections=tuple(sections),
                             placeholders=placeholders or {}, config=config())


def required(*keys) -> dict[str, PlaceholderDefinition]:
    return {k: PlaceholderDefinition(required=True) for k in keys}


class DictAssets:
    def __init__(self, assets: dict[str, Asset]) -> None:
        self._assets = assets

    def get(self, asset_id: str) -> Asset | None:
        return self._assets.get(asset_id)


PNG_1x1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAAC0lEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
)


def resolve(bp, context, assets=None, diag=None):
    diag = diag or DiagnosticCollector()
    assets = assets or DictAssets({})
    return DocumentResolver(context, assets, diag).resolve(bp), diag


# ==========================================================================
# resolver
# ==========================================================================

def test_scalar_placeholder_resolves():
    bp = blueprint([section([paragraph(text_seg("Hi "), ph_seg("org_name"))])],
                   required("org_name"))
    ctx = RenderContext(scalars={"org_name": {"ENG": "Acme LLC"}})

    doc, diag = resolve(bp, ctx)

    runs = doc.sections[0].blocks[0].runs
    assert [r.text for r in runs] == ["Hi ", "Acme LLC"]
    assert diag.warnings == []


def test_missing_required_placeholder_raises():
    bp = blueprint([section([paragraph(ph_seg("org_name"))])], required("org_name"))
    with pytest.raises(PlaceholderError):
        resolve(bp, RenderContext(scalars={}))


def test_missing_optional_placeholder_warns_and_blanks():
    bp = blueprint([section([paragraph(ph_seg("note"))])],
                   {"note": PlaceholderDefinition(required=False)})
    doc, diag = resolve(bp, RenderContext(scalars={}))

    assert doc.sections[0].blocks[0].runs[0].text == ""
    assert len(diag.warnings) == 1


def test_joined_placeholder_joins_present_values():
    seg = JoinedPlaceholderSegment(
        type="placeholder_join",
        items=(ph_seg("first"), ph_seg("last")),
        separator=" ", style=text_style(),
    )
    bp = blueprint([section([paragraph(seg)])], required("first", "last"))
    ctx = RenderContext(scalars={"first": {"ENG": "Ada"}, "last": {"ENG": "Lovelace"}})

    doc, _ = resolve(bp, ctx)
    assert doc.sections[0].blocks[0].runs[0].text == "Ada Lovelace"


def test_image_resolves_when_present_and_warns_when_missing():
    img = ImageSegment(type="image", asset_id="logo", width_emu=990000, height_emu=792000)
    bp = blueprint([section([paragraph(img)])])
    assets = DictAssets({"logo": Asset(data=PNG_1x1, mime="image/png")})

    doc, diag = resolve(bp, RenderContext(), assets)
    run = doc.sections[0].blocks[0].runs[0]
    assert isinstance(run, ResolvedImageRun)
    assert run.width_emu == 990000 and run.asset.mime == "image/png"
    assert diag.warnings == []

    # missing asset -> warn + no run
    doc2, diag2 = resolve(bp, RenderContext(), DictAssets({}))
    assert doc2.sections[0].blocks[0].runs == ()
    assert len(diag2.warnings) == 1


def test_invl_row_expansion():
    table = TableBlueprint(
        type="table",
        rows=(
            RowBlueprint(type="row", cells=(static_cell("Desc"), static_cell("Qty")),
                         style=RowStyleBlueprint(height=0, header=True)),
            RowPlaceholder(type="placeholder_row",
                           cells=(column_cell("invl_desc"), column_cell("invl_qnty")),
                           style=RowStyleBlueprint(height=0, header=False)),
        ),
        style=table_style(),
    )
    bp = blueprint([section([table])])
    ctx = RenderContext(table=InvoiceTableData(
        rows=(
            InvoiceLineRow(values={"invl_desc": {"ENG": "Widget"}, "invl_qnty": {"ENG": "2"}}),
            InvoiceLineRow(values={"invl_desc": {"ENG": "Gadget"}, "invl_qnty": {"ENG": "5"}}),
        ),
        show_tax=False, subtotal={"ENG": "100"}, total_tax=None, total={"ENG": "100"}, labels={},
    ))

    doc, _ = resolve(bp, ctx)
    rtable = doc.sections[0].blocks[0]
    # header + 2 data rows
    assert len(rtable.rows) == 3

    def cell_text(row, col):
        return rtable.rows[row].cells[col].blocks[0].runs[0].text

    assert cell_text(0, 0) == "Desc"
    assert (cell_text(1, 0), cell_text(1, 1)) == ("Widget", "2")
    assert (cell_text(2, 0), cell_text(2, 1)) == ("Gadget", "5")


def test_invoice_table_builds_header_lines_and_summary():
    cols = ["invl_n", "invl_desc", "invl_unit", "invl_qnty", "invl_price", "invl_tax", "invl_total"]
    labels = {c: {"ENG": c} for c in cols}
    labels |= {"subtotal": {"ENG": "Subtotal"}, "total_tax": {"ENG": "Tax"}, "total": {"ENG": "Total"}}
    line = InvoiceLineRow(values={c: {"ENG": f"{c}-v"} for c in cols})

    tp = TablePlaceholder(type="placeholder_table", language="ENG",
                          text_style=text_style(), style=table_style())
    bp = blueprint([section([tp])])
    ctx = RenderContext(table=InvoiceTableData(
        rows=(line,), show_tax=True,
        subtotal={"ENG": "90"}, total_tax={"ENG": "10"}, total={"ENG": "100"}, labels=labels,
    ))

    doc, _ = resolve(bp, ctx)
    rtable = doc.sections[0].blocks[0]
    # header + 1 line + subtotal + tax + total
    assert len(rtable.rows) == 5
    assert len(rtable.rows[0].cells) == 7                     # tax column present
    assert rtable.rows[0].cells[0].blocks[0].runs[0].text == "invl_n"
    assert rtable.rows[0].cells[0].blocks[0].runs[0].style.bold is True     # header bold
    # total row: label + amount, bold
    total_row = rtable.rows[-1]
    assert total_row.cells[-1].blocks[0].runs[0].text == "100"
    assert total_row.cells[-1].blocks[0].runs[0].style.bold is True


def test_invoice_table_drops_tax_column_when_no_tax():
    cols = ["invl_n", "invl_desc", "invl_unit", "invl_qnty", "invl_price", "invl_tax", "invl_total"]
    labels = {c: {"ENG": c} for c in cols} | {"total": {"ENG": "Total"}}
    line = InvoiceLineRow(values={c: {"ENG": "x"} for c in cols})

    tp = TablePlaceholder(type="placeholder_table", language="ENG",
                          text_style=text_style(), style=table_style())
    ctx = RenderContext(table=InvoiceTableData(
        rows=(line,), show_tax=False, subtotal={"ENG": "100"}, total_tax=None,
        total={"ENG": "100"}, labels=labels,
    ))
    doc, _ = resolve(blueprint([section([tp])]), ctx)
    # tax column dropped -> 6 columns
    assert len(doc.sections[0].blocks[0].rows[0].cells) == 6


# ==========================================================================
# docx emitter (round-trip through python-docx)
# ==========================================================================

def emit(bp, context, assets=None):
    diag = DiagnosticCollector()
    resolved = DocumentResolver(context, assets or DictAssets({}), diag).resolve(bp)
    data = DocxEmitter(diag).emit(resolved)
    return data, diag


def open_docx(data: bytes) -> Document:
    return Document(BytesIO(data))


def test_emit_paragraph_opens_and_keeps_text():
    bp = blueprint([section([paragraph(text_seg("Invoice for "), ph_seg("org_name"))])],
                   required("org_name"))
    data, diag = emit(bp, RenderContext(scalars={"org_name": {"ENG": "Acme LLC"}}))

    doc = open_docx(data)                                    # raises if the package is malformed
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Invoice for Acme LLC" in body_text
    assert "{{" not in body_text
    assert diag.warnings == []


def test_emit_table_opens_with_cells():
    table = TableBlueprint(
        type="table",
        rows=(RowBlueprint(type="row", cells=(static_cell("Item"), static_cell("Total")),
                           style=RowStyleBlueprint(height=0, header=True)),),
        style=table_style(),
    )
    data, _ = emit(blueprint([section([table])]), RenderContext())

    doc = open_docx(data)
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "Item"
    assert doc.tables[0].rows[0].cells[1].text == "Total"


def test_emit_default_header():
    header = HeaderFooterGroupBlueprint(
        default=HeaderFooterBlueprint(type=HeaderFooterType.DEFAULT,
                                      blocks=(paragraph(text_seg("ACME header")),)),
        first=None, even=None,
    )
    sec = section([paragraph(text_seg("body"))], headers=header)
    data, _ = emit(blueprint([sec]), RenderContext())

    doc = open_docx(data)
    header_text = "\n".join(p.text for p in doc.sections[0].header.paragraphs)
    assert "ACME header" in header_text


def test_emit_embeds_image_media_part():
    img = ImageSegment(type="image", asset_id="logo", width_emu=990000, height_emu=792000)
    bp = blueprint([section([paragraph(img)])])
    assets = DictAssets({"logo": Asset(data=PNG_1x1, mime="image/png")})
    data, diag = emit(bp, RenderContext(), assets)

    open_docx(data)                                          # valid package
    names = zipfile.ZipFile(BytesIO(data)).namelist()
    assert any(n.startswith("word/media/") for n in names)
    assert diag.warnings == []


def test_emit_multi_section_produces_two_sections():
    s1 = section([paragraph(text_seg("section one"))])
    s2 = section([paragraph(text_seg("section two"))])
    data, diag = emit(blueprint([s1, s2]), RenderContext())

    doc = open_docx(data)
    assert len(doc.sections) == 2
    assert "flattened" not in " ".join(w.message for w in diag.warnings)


# ==========================================================================
# end to end: ingest -> resolve -> emit
# ==========================================================================

def test_ingest_resolve_emit_round_trip(make_docx, fixture_provider):
    from app.document_engine.orchestration.pipeline import TemplateIngestionPipeline

    path = make_docx(paragraphs=["Invoice for {{ org_name }}", "No {{ invoice_no }}"])
    pipeline = TemplateIngestionPipeline(fixture_provider)
    bp = pipeline.finalize(pipeline.ingest(path).draft)

    ctx = RenderContext(scalars={
        "org_name": {"ENG": "Acme LLC"},
        "invoice_no": {"ENG": "INV-001"},
    })
    data, diag = emit(bp, ctx)

    doc = open_docx(data)
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Acme LLC" in body_text and "INV-001" in body_text
    assert "{{" not in body_text
